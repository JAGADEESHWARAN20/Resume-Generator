
import os
import tempfile
import pdf2image
import PyPDF4
import tkinter as tk
from tkinter import ttk
from pdf2image.pdf2image import convert_from_path
from io import BytesIO
import tkinter as tk
from tkinter import * #type:ignore
from tkinter import StringVar, ttk
from tkinter import filedialog
import ttkbootstrap as tb
import sqlite3
from tkinter.simpledialog import askstring
from ttkbootstrap import Style
from tkinter import messagebox
import datetime
from docxtpl import DocxTemplate
from pathlib import Path
from docx2pdf import convert
from PIL import Image, ImageTk
from io import BytesIO
import docx
from docx import Document
from docx.shared import Cm
from tkPDFViewer import tkPDFViewer as pdf
from PIL import ImageTk as IMAGETEXT
from tkinter import filedialog
import subprocess
import fitz





root = tb.Window(themename="cosmo")
root.title('Student Login')
root.geometry('800x900')
save_first_time= True
width_of_window = 1120
height_of_window = 1000

screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

x_coordinate = (screen_width / 2) - (width_of_window / 2)
y_coordinate = (screen_height / 2) - (height_of_window / 1.85)

root.geometry("%dx%d+%d+%d" % (width_of_window, height_of_window, x_coordinate, y_coordinate))



def on_close():
    root.destroy()

root.protocol("WM_DELETE_WINDOW", on_close)

# Set the ttkbootstrap style
style = Style(theme='vapor')

# Create a parent frame with a 5x5 grid layout
parent_frame = tb.Frame(root)
parent_frame.pack(fill='both', expand=True)
parent_frame.columnconfigure(1, weight=1)
parent_frame.rowconfigure(2, weight=1)

Nav_bar = tb.LabelFrame(parent_frame,text='Navbar',width=200,height=900)
Nav_bar.pack(fill='y',side='left',padx=10,pady=10)

notification_bar = tb.LabelFrame(parent_frame,text='Notification',width=200,height=900)
notification_bar.pack(fill='y',side='right',padx=10,pady=10)

table = ttk.Treeview(notification_bar, columns=('Date', 'Message'), show='headings')
table.column('Date', width=80, anchor='e')
table.column('Message', width=150, anchor='w')
table.heading('Date', text='Date')
table.heading('Message', text='Message')
table.pack(side='left',fill='y',padx=10,pady=10)



mainframe = tb.Frame(parent_frame,width=700,height=800)
mainframe.pack(fill='x',padx=10,pady=10)

main_label_frame = tb.LabelFrame(mainframe,text='Main Frame',width=700,height=800)
main_label_frame.pack(side='top',fill='x',expand=True,ipadx=300,ipady=500)

tb.Label(main_label_frame,bootstyle='secondary',text="Welcome To Resume Generator",font=("Arial", 19)).pack(side=TOP,padx=20,pady=20,anchor='nw') #type:ignore
tb.Label(main_label_frame,bootstyle='info',text="Place to Prepare New Resume",font=("Arial", 19)).pack(side=TOP,padx=20,pady=20,anchor='nw') #type:ignore


conn = sqlite3.connect('MainDB.db')
            # Create a cursor object to execute SQL commands
c = conn.cursor()

# Fetch all records from the table
c.execute('''SELECT date,message FROM Records''')
records = c.fetchall()

# Insert the records into the table widget
for record in records:
    table.insert('', 'end', values=record)

# Save the changes
conn.commit()
messagebox.showinfo("Success", "Your Data Fetched Successfully!")
conn.close()







def logout():
    subprocess.Popen("RegistrationWindow.exe")
    root.destroy()
    



#Bottom Buttons
Log_out = tb.Button(Nav_bar,text='LOG OUT',bootstyle='primary',command=logout) #type:ignore
Log_out.pack(side='bottom',pady=10,ipadx=20,padx=10)

style = Style(theme='vapor')

def Mode_Change():
    current_theme = Style().theme_use()
    if current_theme == 'vapor':
        Style(theme='morph').theme_use()
        tb.Style = 'morph'
    else:
        Style(theme='vapor').theme_use()
        tb.Style = 'vapor'

    
        
Mode_change = tb.Checkbutton(Nav_bar,text='MODE',bootstyle='info-round-toggle',command=Mode_Change)  #type:ignore
Mode_change.pack(side='bottom',pady=10,padx=10)


def renderimage():

    filename = filedialog.askopenfilename()

    with open(filename, 'rb') as f:
        img_data = f.read()

    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS images
    (id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT,
    data BLOB)''')
    c.execute("INSERT INTO images (name, data) VALUES (?, ?)", (filename, img_data))
    conn.commit()

    doc = docx.Document(Path(__file__).parent / "Resumetmpl.docx")
    output_path = Path(__file__).parent / "Output_resume.docx"

    placeholder = '{{image}}'

    for paragraph in doc.paragraphs:
        # Check if the paragraph contains the placeholder text
        if placeholder in paragraph.text:
            # Remove the placeholder text
            text = paragraph.text.replace(placeholder, '').strip()
            # Add the image
            paragraph.clear()
            run = paragraph.add_run()
            c.execute("SELECT data FROM images WHERE name=?", (filename,))
            img_data = c.fetchone()[0]
            img = Image.open(BytesIO(img_data))
            img = img.resize((int(img.width*0.75), int(img.height*0.85)))
            img_path = Path(filename).stem + '.png'
            img.save(img_path)
            run.add_picture(img_path, width=Cm(5), height=Cm(6.5))
            # Add the original text after the image
            run.add_text(' ' + text)

    doc.save(output_path)
    messagebox.showinfo('Success','Your Image is Inserted Sucessfully')
    conn.close()

def Downloadcv():
    global rrn_value
    conn = sqlite3.connect('MainDB.db')
    c = conn.cursor()
    template_path = Path(__file__).parent / "Output_resume.docx"
    output_path = Path(__file__).parent / f"Resume{rrn_value}.docx"
    
    doc = DocxTemplate(template_path)

    c.execute("SELECT name, dob, mobile, address, email, linkedin, website, company_name_one, duration_one, job_title_one, skills_learned_one, company_name_two, duration_two, job_title_two, skills_learned_two, company_name_three, duration_three, job_title_three, skills_learned_three, project_name_one, explain_one_project, project_name_two, explain_two_project, project_name_three, explain_three_project, university_name, degree, stream, year_university, cgpa_university, high_school_name, branch, year_high_school, cgpa_high_school, elementary_school_name, year_elementary_school, cgpa_elementary_school, skills FROM Resume WHERE rrn_value = ?", (rrn_value,))
    row = c.fetchone()
    
    
    context = {
    "Name": row[0], 
    "DOB": row[1], 
    "Mobile": row[2], 
    "Address": row[3], 
    "Email": row[4], 
    "Linkedin": row[5], 
    "Website": row[6],    
    "Company_name_one": row[7], 
    "Duration_one": row[8], 
    "Job_Title_one": row[9], 
    "Skills_learned_one": row[10], 
    "Company_name_two": row[11], 
    "Duration_two": row[12], 
    "Job_Title_two": row[13], 
    "Skills_learned_two": row[14],           
    "Company_name_three": row[15], 
    "Duration_three": row[16], 
    "Job_Title_three": row[17], 
    "Skills_learned_three": row[18],
    "project_Name_one": row[19],  
    "Explain_one": row[20], 
    "project_Name_two": row[21], 
    "Explain_two": row[22], 
    "project_Name_three": row[23], 
    "Explain_three": row[24], 
    "University": row[25], 
    "degree": row[26], 
    "stream": row[27], 
    "Year_one": row[28], 
    "cgpa": row[29],
    "High_school_name": row[30], 
    "branch": row[31], 
    "Year_two": row[32], 
    "cgpa_two": row[33],
    "Elementry_school_name": row[34], 
    "Year_three": row[35], 
    "cgpa_three": row[36], 
    "skills_two": row[37]
    }


    doc.render(context)
    doc.save(output_path)

    
    conn.close()
    convert(f"Resume{rrn_value}.docx",f"Resume{rrn_value}.pdf")

    
    messagebox.showinfo("Success", "Resume is ready!")




def Personal_frame():
    global rrn_value
    global updater_frame
    Style(theme='cyborg')
    if updater_frame is not None:
        updater_frame.config(bootstyle='warning') #type:ignore
        # destroy any existing frames inside updater_frame
        for child in updater_frame.winfo_children():
            child.destroy()
            
        # create and add the Personal_details frame
        Personal_details = ttk.LabelFrame(updater_frame, text='Personal Details', width=700, height=800,bootstyle='warning') #type:ignore
        Personal_details.pack(fill='both', padx=10, pady=10,ipady=300)
        
        conn = sqlite3.connect('MainDB.db')

        # Create a cursor object to execute SQL commands
        c = conn.cursor()

        # Create a table to store experience information
        c.execute('''CREATE TABLE IF NOT EXISTS Resume (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    rrn_value INTEGER UNIQUE,
                    name TEXT,
                    DOB DATE,
                    mobile VARCHAR(20),
                    address VARCHAR(255),
                    email VARCHAR(255),
                    linkedin VARCHAR(255),
                    website VARCHAR(255),
                    company_name_one VARCHAR(255),
                    duration_one VARCHAR(20),
                    job_title_one VARCHAR(255),
                    skills_learned_one VARCHAR(255),
                    company_name_two VARCHAR(255),
                    duration_two VARCHAR(20),
                    job_title_two VARCHAR(255),
                    skills_learned_two VARCHAR(255),
                    company_name_three VARCHAR(255),
                    duration_three VARCHAR(20),
                    job_title_three VARCHAR(255),
                    skills_learned_three VARCHAR(255),
                    project_name_one VARCHAR(255),
                    explain_one_project TEXT,
                    project_name_two VARCHAR(255),
                    explain_two_project TEXT,
                    project_name_three VARCHAR(255),
                    explain_three_project TEXT,
                    university_name VARCHAR(255),
                    degree TEXT,
                    stream TEXT,
                    year_university VARCHAR(4),
                    cgpa_university FLOAT,
                    high_school_name VARCHAR(255),
                    branch TEXT,
                    year_high_school VARCHAR(4),
                    cgpa_high_school FLOAT,
                    elementary_school_name VARCHAR(255),
                    year_elementary_school VARCHAR(4),
                    cgpa_elementary_school FLOAT,
                    skills VARCHAR(255)
                );''')


        # Save the changes
        conn.commit()
        conn.close()

        
        rrn = tb.Label(Personal_details,text='RRN',bootstyle='warning') #type:ignore
        rrn.pack(side='top',anchor='nw',padx=10,pady=10)
        rrn_entry = tb.Label(Personal_details,bootstyle='warning',text=rrn_value,font=('poppins',18)) #type:ignore
        rrn_entry.pack(side='top',anchor='nw',padx=10,pady=10)

        Name = tb.Label(Personal_details,text='Name')
        Name.pack(side='top',anchor='nw',padx=10,pady=10)
        Name_entry = tb.Entry(Personal_details)
        Name_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
        
        image_label_frame = ttk.LabelFrame(Personal_details,text='Select Profile')
        image_label_frame.pack(side='top',anchor='nw',padx=10,pady=10,ipadx=20,ipady=20,fill=X)
        
        select_image = ttk.Button(image_label_frame,text='Select & Insert',command=renderimage)
        select_image.pack(side=LEFT,anchor='center',ipadx=100,padx=20,pady=20)
        
        Dob = tb.Label(Personal_details,text='Dob')
        Dob.pack(side='top',anchor='nw',padx=10,pady=10)
        Dob_entry = tb.Entry(Personal_details)
        Dob_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        Address = tb.Label(Personal_details,text='Address')
        Address.pack(side='top',anchor='nw',padx=10,pady=10)
        Address_entry = tb.Entry(Personal_details)
        Address_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
        
        mobile = tb.Label(Personal_details,text='Mobile')
        mobile.pack(side='top',anchor='nw',padx=10,pady=10)
        mobile_entry = tb.Entry(Personal_details)
        mobile_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        email = tb.Label(Personal_details,text='Email')
        email.pack(side='top',anchor='nw',padx=10,pady=10)
        email_entry = tb.Entry(Personal_details)
        email_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')

        Linkedin = tb.Label(Personal_details,text='Linkedin')
        Linkedin.pack(side='top',anchor='nw',padx=10,pady=10)
        Linkedin_entry = tb.Entry(Personal_details)
        Linkedin_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')

        website = tb.Label(Personal_details,text='website')
        website.pack(side='top',anchor='nw',padx=10,pady=10)
        website_entry = tb.Entry(Personal_details)
        website_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
        def insertdata():
            global rrn_value
            name = Name_entry.get()
            dob = Dob_entry.get()
            address = Address_entry.get()
            mobile = mobile_entry.get()
            email = email_entry.get()
            linkedin = Linkedin_entry.get()
            website = website_entry.get()

            conn = sqlite3.connect('MainDB.db')
            # Create a cursor object to execute SQL commands
            c = conn.cursor()

            c.execute('''INSERT INTO Resume (rrn_value, name, DOB, mobile, address, email, linkedin, website)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                        (rrn_value, name, dob, mobile, address, email, linkedin, website))

            current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            frame_name = "Personal Information"
            message = "Changes made in " + frame_name

            # Create a new table to hold the records if it doesn't exist
            c.execute('''CREATE TABLE IF NOT EXISTS Records
                        (id INTEGER PRIMARY KEY AUTOINCREMENT,
                        date TEXT,
                        message TEXT)''')

            # Insert a new record into the table
            c.execute('''INSERT INTO Records (date, message)
                        VALUES (?, ?)''', (current_time, message))

            # Fetch all records from the table
            c.execute('''SELECT date,message FROM Records''')
            records = c.fetchall()

            # Insert the records into the table widget
            for record in records:
                table.insert('', 'end', values=record)

            # Save the changes
            conn.commit()
            messagebox.showinfo("Success", "Your Data Inserted Successfully!")
            conn.close()

            

        insert_data = tb.Button(Personal_details,bootstyle='primary',text='Insert Data',command=insertdata) #type:ignore
        insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  
        
        Next_button = tb.Button(Personal_details,bootstyle='primary',text='Next',command=Education_frame) #type:ignore
        Next_button.pack(side='right',anchor='se',pady=5,padx=10)
    
        save_button = tb.Button(Personal_details,bootstyle='primary',text='Save') #type:ignore
        save_button.pack(side='right',anchor='se',pady=5,padx=10)


        Download_cv = tb.Button(Personal_details,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
        Download_cv.pack(side='bottom',anchor='s',pady=5,padx=10)      
        
    else:
        # if updater_frame does not exist, create it and add the Personal_details frame
        Updater_frame()
        Personal_details = ttk.LabelFrame(updater_frame, text='Personal Details', width=700, height=800)
        Personal_details.pack(fill='both', padx=10, pady=10)


def Home():
    for child in main_label_frame.winfo_children():
        child.destroy()
    main_label_frame.config(text='Home')
    tb.Label(main_label_frame,bootstyle='secondary',text="Welcome To Resume Generator",font=("Arial", 19)).pack(side=TOP,padx=20,pady=20,anchor='nw') #type:ignore
    tb.Label(main_label_frame,bootstyle='info',text="Place to Prepare New Resume",font=("Arial", 19)).pack(side=TOP,padx=20,pady=20,anchor='nw') #type:ignore



updater_frame = None

def Updater_frame():
    global updater_frame
    

    # access the global variable
    Style(theme='vapor')
    
    if updater_frame is None:
        for child in main_label_frame.winfo_children():
            child.destroy()

        global rrn_value
        rrn_value = askstring("  ","Enter RRN") # type: ignore

        updater_frame = ttk.LabelFrame(main_label_frame, text='Updater Frame', width=700, height=800)
        updater_frame.pack(fill='both', padx=10, pady=10)
        
        Personal_frame()

    elif updater_frame is not None:
        for child in main_label_frame.winfo_children():
            child.destroy()
        
        updater_frame = ttk.LabelFrame(main_label_frame, text='Updater Frame', width=700, height=800)
        updater_frame.pack(fill='both', padx=10, pady=10)

    else:
        updater_frame.destroy()
        updater_frame = None



def DashboardFrame(pdf_location):
    # Clear any existing widgets in the mainframe
    for child in mainframe.winfo_children():
        child.destroy()
        
    # Create a new LabelFrame for the Dashboard
    dashboard_frame = ttk.LabelFrame(mainframe, text='Dashboard', width=200, height=800)
    dashboard_frame.pack(side='top', fill='both', expand=True)

    # Create a canvas to display the PDF
    canvas = tk.Canvas(dashboard_frame, bg='white', width=200, height=800)
    canvas.pack(side='left', fill='both', expand=True)

    # Add a scrollbar to the canvas
    scrollbar = ttk.Scrollbar(dashboard_frame, style='primary.Vertical.TScrollbar', orient='vertical', command=canvas.yview) #type:ignore
    scrollbar.pack(side='right', fill='y')

    # Associate the scrollbar with the canvas
    canvas.config(yscrollcommand=scrollbar.set)
    canvas.config(scrollregion=canvas.bbox('all'))

    # Load the PDF and display it in the canvas
    v1 = pdf.ShowPdf()
    pdf_location = open(f"Resume{rrn_value}.pdf","r").name
    v2 = v1.pdf_view(canvas, pdf_location=pdf_location, width=700, height=800)
    canvas.create_window((0, 0), anchor='nw', window=v2)    

    # Update the canvas when the PDF is updated 
    canvas.update()



def Experience():
    global updater_frame
   
    Style(theme='vapor')
    
    if updater_frame is not None:
        updater_frame.config(bootstyle='light') #type:ignore
        # destroy any existing frames inside updater_frame
        for child in updater_frame.winfo_children():
            child.destroy()
        
        
        frame = ttk.Frame(updater_frame,width=600,height=800)
        frame.pack(fill='both',side='left', expand=True,anchor='center')


        labelframe1 = ttk.Labelframe(frame, text='Job Experience 1',width=600,height=800)
        labelframe1.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
        
        
        def FrameJobEXperience2():
            if updater_frame is not None:
                for child in updater_frame.winfo_children():
                    child.destroy()
                    
            frame = ttk.Frame(updater_frame,width=600,height=800)
            frame.pack(fill='both',side='left', expand=True,anchor='center')
            labelframe2 = ttk.Labelframe(frame, text='Job Experience 2',width=600,height=800)
            labelframe2.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
            #labelframe2
            ComapnyName = tb.Label(labelframe2,text='Comapny Name')
            ComapnyName.pack(side='top',anchor='nw',padx=10,pady=10)
            ComapnyName_entry = tb.Entry(labelframe2)
            ComapnyName_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
            
            job_Title = tb.Label(labelframe2,text='Job Title')
            job_Title.pack(side='top',anchor='nw',padx=10,pady=10)
            job_Title_entry = tb.Entry(labelframe2)
            job_Title_entry.pack(side='top',anchor='nw',padx=10,pady=10)
            
            skillsLearned = tb.Label(labelframe2,text='skills Learned')
            skillsLearned.pack(side='top',anchor='nw',padx=10,pady=10)
            skillsLearned_entry = tb.Entry(labelframe2)
            skillsLearned_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
            
            Duration = tb.Label(labelframe2,text='Duration')
            Duration.pack(side='top',anchor='nw',padx=10,pady=10)
            Duration_entry = tb.Entry(labelframe2)
            Duration_entry.pack(side='top',anchor='nw',padx=10,pady=10)
            def insertData():
                global rrn_value
                
                company_name2 = ComapnyName_entry.get()
                job_title2  = job_Title_entry.get()
                skills_learned2  = skillsLearned_entry.get()
                duration2  = Duration_entry.get()
                conn = sqlite3.connect('MainDB.db')

                # Create a cursor object to execute SQL commands
                c = conn.cursor()
            
                try:
                    print(rrn_value,company_name2, job_title2, skills_learned2, duration2)
                    c.execute('''UPDATE Resume
                                SET company_name_two = ?,
                                    duration_two = ?,
                                    job_title_two = ?,
                                    skills_learned_two = ?
                                WHERE rrn_value = ?''', 
                                (company_name2, duration2, job_title2, skills_learned2, rrn_value))
                        
                    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                    frame_name = "Experience Frame"
                    message = {rrn_value}, "Changes made in " + frame_name

                    # Insert the records into the table widget
                    table.insert('', 'end', values=(current_time,message))

                    conn.commit()
                    messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                except Exception as e:
                    messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                conn.close()
                                                
            insert_data = tb.Button(labelframe2,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
            insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  
            
            save_button = tb.Button(labelframe2,bootstyle='primary',text='Save') #type:ignore
            save_button.pack(side='left',anchor='se',pady=5,padx=10)
            
            Add_button = tb.Button(labelframe2,bootstyle='primary',text='ADD',command=FrameJobEXperience3)  #type:ignore
            Add_button.pack(side='left',anchor='sw',pady=5,padx=10) 
          
            Download_cv = tb.Button(labelframe2,bootstyle='primary',text='Download CV') #type:ignore
            Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
        
            Next_button = tb.Button(labelframe2,bootstyle='primary',text='Next',command=Skills_frame) #type:ignore
            Next_button.pack(side='right',anchor='se',pady=5,padx=10)
        
        def FrameJobEXperience3():
            if updater_frame is not None:
                for child in updater_frame.winfo_children():
                    child.destroy()
                    
            frame = ttk.Frame(updater_frame,width=600,height=800)
            frame.pack(fill='both',side='left', expand=True,anchor='center')       
             
            labelframe3 = ttk.Labelframe(frame, text='Job Experience 3',width=600,height=800)
            labelframe3.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
            
            #labelframe3
            ComapnyName = tb.Label(labelframe3,text='Comapny Name')
            ComapnyName.pack(side='top',anchor='nw',padx=10,pady=10)
            ComapnyName_entry = tb.Entry(labelframe3)
            ComapnyName_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
            
            job_Title = tb.Label(labelframe3,text='Job Title')
            job_Title.pack(side='top',anchor='nw',padx=10,pady=10)
            job_Title_entry = tb.Entry(labelframe3)
            job_Title_entry.pack(side='top',anchor='nw',padx=10,pady=10)
            
            skillsLearned = tb.Label(labelframe3,text='skills Learned')
            skillsLearned.pack(side='top',anchor='nw',padx=10,pady=10)
            skillsLearned_entry = tb.Entry(labelframe3)
            skillsLearned_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
            
            Duration = tb.Label(labelframe3,text='Duration')
            Duration.pack(side='top',anchor='nw',padx=10,pady=10)
            Duration_entry = tb.Entry(labelframe3)
            Duration_entry.pack(side='top',anchor='nw',padx=10,pady=10)
            
            def insertData():
                global rrn_value
                company_name3 = ComapnyName_entry.get()
                job_title3  = job_Title_entry.get()
                skills_learned3  = skillsLearned_entry.get()
                duration3  = Duration_entry.get()
                conn = sqlite3.connect('MainDB.db')

                # Create a cursor object to execute SQL commands
                c = conn.cursor()
                try:
                    print(rrn_value,company_name3, job_title3, skills_learned3, duration3)
                    c.execute('''UPDATE Resume
                        SET company_name_three = ?,
                            job_title_three = ?,
                            skills_learned_three = ?,
                            duration_three = ?
                        WHERE rrn_value = ?''', (company_name3, job_title3, skills_learned3, duration3, rrn_value))

                    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    
                    frame_name = "Experience Frame"
                    message = {rrn_value}, "Changes made in " + frame_name

                    table.insert('', 'end', values=(current_time,message))

                    conn.commit()
                    messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                except Exception as e:
                    messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                    conn.close()

            insert_data = tb.Button(labelframe3,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
            insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  
            
            save_button = tb.Button(labelframe3,bootstyle='primary',text='Save') #type:ignore
            save_button.pack(side='left',anchor='se',pady=5,padx=10)
            
         
            Download_cv = tb.Button(labelframe3,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
            Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
        
            Next_button = tb.Button(labelframe3,bootstyle='primary',text='Next',command=Skills_frame) #type:ignore
            Next_button.pack(side='right',anchor='se',pady=5,padx=10)
    

        #labelframe1            
        ComapnyName = tb.Label(labelframe1,text='Comapny Name')
        ComapnyName.pack(side='top',anchor='nw',padx=10,pady=10)
        ComapnyName_entry = tb.Entry(labelframe1)
        ComapnyName_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
        
        job_Title = tb.Label(labelframe1,text='Job Title')
        job_Title.pack(side='top',anchor='nw',padx=10,pady=10)
        job_Title_entry = tb.Entry(labelframe1)
        job_Title_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        skillsLearned = tb.Label(labelframe1,text='skills Learned')
        skillsLearned.pack(side='top',anchor='nw',padx=10,pady=10)
        skillsLearned_entry = tb.Entry(labelframe1)
        skillsLearned_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
        
        Duration = tb.Label(labelframe1,text='Duration')
        Duration.pack(side='top',anchor='nw',padx=10,pady=10)
        Duration_entry = tb.Entry(labelframe1)
        Duration_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        def insertData():
            global rrn_value
            
            company_name = ComapnyName_entry.get()
            job_title  = job_Title_entry.get()
            skills_learned  = skillsLearned_entry.get()
            duration  = Duration_entry.get()
            
            conn = sqlite3.connect('MainDB.db')

            # Create a cursor object to execute SQL commands
            c = conn.cursor()
            try:
                print(rrn_value,company_name, job_title, skills_learned, duration)
                c.execute(f'''UPDATE Resume
                                SET company_name_one = '{company_name}',
                                    job_title_one = '{job_title}',
                                    skills_learned_one = '{skills_learned}',
                                    duration_one = '{duration}'
                                WHERE rrn_value = {rrn_value};
                                ''')
                current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
                frame_name = "Experience Frame"
                message = {rrn_value}, "Changes made in " + frame_name

                # Insert the records into the table widget
                table.insert('', 'end', values=(current_time,message))



                conn.commit()
                messagebox.showinfo("Success", "Your Data Inserted Successfully!")
            except Exception as e:
                messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                conn.close()

        insert_data = tb.Button(labelframe1,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
        insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

        save_button = tb.Button(labelframe1,bootstyle='primary',text='Save') #type:ignore
        save_button.pack(side='left',anchor='se',pady=5,padx=10)
        
        Add_button = tb.Button(labelframe1,bootstyle='primary',text='ADD',command=FrameJobEXperience2)  #type:ignore
        Add_button.pack(side='left',anchor='sw',pady=5,padx=10)
        

        Download_cv = tb.Button(labelframe1,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
        Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
    
        Next_button = tb.Button(labelframe1,bootstyle='primary',text='Next',command=Skills_frame) #type:ignore
        Next_button.pack(side='right',anchor='se',pady=5,padx=10)

    else:
        # if updater_frame does not exist, create it and add the Experience_frame
        Updater_frame()
        Experience_frame = ttk.LabelFrame(updater_frame, text='Experience_frame', width=700, height=800)
        Experience_frame.pack(fill='both', padx=10, pady=10)


def Education_frame():
    global updater_frame
    
    if updater_frame is not None:
        updater_frame.config(bootstyle='primary')#type:ignore
        for child in updater_frame.winfo_children():
            child.destroy()
        

        frame = ttk.Frame(updater_frame,width=600,height=800)
        frame.pack(fill='both',side='left', expand=True,anchor='center')
        


        labelframe1 = ttk.Labelframe(frame, text='Graduation',width=600,height=450)
        labelframe1.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
        
        def frame2():
            for child in frame.winfo_children():
                child.destroy()
            labelframe2 = ttk.Labelframe(frame, text='Higher education',width=600,height=420)
            labelframe2.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
            
            #labelframe1            
            University = tb.Label(labelframe2,text='School Name')
            University.pack(side='top',anchor='nw',padx=10,pady=10)
            University_entry2 = tb.Entry(labelframe2)
            University_entry2.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
            
            branch = tb.Label(labelframe2,text='branch')
            branch.pack(side='top',anchor='nw',padx=10,pady=10)
            branch_entry2 = tb.Entry(labelframe2)
            branch_entry2.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
      
            Cgpa = tb.Label(labelframe2,text='CGPA/Percentage')
            Cgpa.pack(side='top',anchor='nw',padx=10,pady=10)
            CGPA_entry2 = tb.Entry(labelframe2)
            CGPA_entry2.pack(side='top',anchor='nw',padx=10,pady=10)
            
            Year = tb.Label(labelframe2,text='Year')
            Year.pack(side='top',anchor='nw',padx=10,pady=10)
            Year_entry2 = tb.Entry(labelframe2)
            Year_entry2.pack(side='top',anchor='nw',padx=10,pady=10)
            
            def insertData():
                global rrn_value
                
                University_name = University_entry2.get()
                branch  = branch_entry2.get()
                Cgpa  = CGPA_entry2.get()
                year = Year_entry2.get()
                
                conn = sqlite3.connect('MainDB.db')

                # Create a cursor object to execute SQL commands
                c = conn.cursor()
                try:
                    print(rrn_value,University_name, branch, year, Cgpa)
                    c.execute('''UPDATE Resume
                                SET high_school_name  = ?,
                                    branch = ?,
                                    year_high_school   = ?,
                                    cgpa_high_school  = ?
                                WHERE rrn_value = ?''',
                                (University_name, branch, year, Cgpa, rrn_value))


                    conn.commit()
                    messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                    
                    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
                    frame_name = "Education Frame"
                    message = {rrn_value}, "Changes made in " + frame_name

                    # Insert the records into the table widget
                    table.insert('', 'end', values=(current_time,message))


                except Exception as e:
                    messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                    conn.close()
                    
            insert_data = tb.Button(labelframe2,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
            insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  
 
            save_button = tb.Button(labelframe2,bootstyle='primary',text='Save') #type:ignore
            save_button.pack(side='left',anchor='se',pady=5,padx=10)

            Download_cv = tb.Button(labelframe2,bootstyle='primary',text='Download CV') #type:ignore
            Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
            
            Add_button = tb.Button(labelframe2,bootstyle='primary',text='ADD',command=frame3) #type:ignore
            Add_button.pack(side='left',anchor='se',pady=5,padx=10)

        
        
            Next_button = tb.Button(labelframe2,bootstyle='primary',text='Next',command=Experience) #type:ignore
            Next_button.pack(side='right',anchor='se',pady=5,padx=10)
            
            
        def frame3():
            for child in frame.winfo_children():
                child.destroy()
            labelframe3 = ttk.Labelframe(frame, text='SSLC',width=600,height=450)
            labelframe3.pack(fill='both', padx=10, pady=10,side='top', expand=True,anchor='center',ipadx=500,ipady=500)
            #labelframe1            
            University3 = tb.Label(labelframe3,text='School Name')
            University3.pack(side='top',anchor='nw',padx=10,pady=10)
            University_entry3 = tb.Entry(labelframe3)
            University_entry3.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')

            Cgpa = tb.Label(labelframe3,text='CGPA/Percentage')
            Cgpa.pack(side='top',anchor='nw',padx=10,pady=10)
            CGPA_entry3 = tb.Entry(labelframe3)
            CGPA_entry3.pack(side='top',anchor='nw',padx=10,pady=10)
            
            Year = tb.Label(labelframe3,text='Year')
            Year.pack(side='top',anchor='nw',padx=10,pady=10)
            Year_entry3 = tb.Entry(labelframe3)
            Year_entry3.pack(side='top',anchor='nw',padx=10,pady=10)
            
            
            def insertData():
                global rrn_value
                
                University_name = University_entry3.get()
                Cgpa  = CGPA_entry3.get()
                year = Year_entry3.get()
                
                conn = sqlite3.connect('MainDB.db')

                # Create a cursor object to execute SQL commands
                c = conn.cursor()
                try:
                    print(rrn_value,University_name, year, Cgpa)
                    c.execute('''UPDATE Resume
                                SET elementary_school_name = ?,
                                    year_elementary_school = ?,
                                    cgpa_elementary_school = ?
                                WHERE rrn_value = ?''',
                                (University_name, year, Cgpa, rrn_value))
                    current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                    frame_name = "Education Frame"
                    message = {rrn_value}, "Changes made in " + frame_name


                    # Insert the records into the table widget
                    table.insert('', 'end', values=(current_time,message))

                    conn.commit()
                    messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                except Exception as e:
                    messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                    conn.close()
                
                
            insert_data = tb.Button(labelframe3,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
            insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  
            

            save_button = tb.Button(labelframe3,bootstyle='primary',text='Save') #type:ignore
            save_button.pack(side='left',anchor='se',pady=5,padx=10)
            
        
            Download_cv = tb.Button(labelframe3,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
            Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
            
        
            Next_button = tb.Button(labelframe3,bootstyle='primary',text='Next',command=Experience) #type:ignore
            Next_button.pack(side='right',anchor='se',pady=5,padx=10)

            
  
        
        #labelframe1            
        University = tb.Label(labelframe1,text='University/Institution')
        University.pack(side='top',anchor='nw',padx=10,pady=10)
        University_entry = tb.Entry(labelframe1)
        University_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='x')
        
        Degree_Title = tb.Label(labelframe1,text='Degree')
        Degree_Title.pack(side='top',anchor='nw',padx=10,pady=10)
        Degree_Title_entry = tb.Entry(labelframe1)
        Degree_Title_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        Stream = tb.Label(labelframe1,text='Stream')
        Stream.pack(side='top',anchor='nw',padx=10,pady=10)
        Stream_entry = tb.Entry(labelframe1)
        Stream_entry.pack(side='top',anchor='nw',padx=10,pady=10,fill='both')
        
        Cgpa = tb.Label(labelframe1,text='CGPA')
        Cgpa.pack(side='top',anchor='nw',padx=10,pady=10)
        CGPA_entry = tb.Entry(labelframe1)
        CGPA_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        Year = tb.Label(labelframe1,text='Year')
        Year.pack(side='top',anchor='nw',padx=10,pady=10)
        Year_entry = tb.Entry(labelframe1)
        Year_entry.pack(side='top',anchor='nw',padx=10,pady=10)
        
        def insertData():
            global rrn_value
            
            University_name = University_entry.get()
            Degree_title  = Degree_Title_entry.get()
            Stream  = Stream_entry.get()
            Cgpa  = CGPA_entry.get()
            year = Year_entry.get()
            
            conn = sqlite3.connect('MainDB.db')

            # Create a cursor object to execute SQL commands
            c = conn.cursor()
            try:
                print(rrn_value,University_name, Degree_title, Stream, year, Cgpa)
                c.execute('''UPDATE Resume
                            SET university_name = ?,
                                degree = ?,
                                stream = ?,
                                year_university = ?,
                                cgpa_university = ?
                            WHERE rrn_value = ?''',
                            (University_name, Degree_title, Stream, year, Cgpa, rrn_value))
                current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                
                frame_name = "Education Frame"
                message = {rrn_value}, "Changes made in " + frame_name
                # Insert the records into the table widget
                table.insert('', 'end', values=(current_time,message))


                conn.commit()
                messagebox.showinfo("Success", "Your Data Inserted Successfully!")
            except Exception as e:
                messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                conn.close()
            
                
        
        insert_data = tb.Button(labelframe1,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
        insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

        save_button = tb.Button(labelframe1,bootstyle='primary',text='Save') #type:ignore
        save_button.pack(side='left',anchor='se',pady=5,padx=10)

        Download_cv = tb.Button(labelframe1,bootstyle='primary',text='Download CV') #type:ignore
        Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
      
        Add_button = tb.Button(labelframe1,bootstyle='primary',text='ADD',command=frame2) #type:ignore
        Add_button.pack(side='left',anchor='se',pady=5,padx=10)

        Next_button = tb.Button(labelframe1,bootstyle='primary',text='Next',command=Experience) #type:ignore
        Next_button.pack(side='right',anchor='se',pady=5,padx=10)
        
        
    else:
        # if updater_frame does not exist, create it and add the Education_frame
        Updater_frame()
        Education_frame = ttk.LabelFrame(updater_frame, text='Education_frame', width=700, height=800)
        Education_frame.pack(fill='both', padx=10, pady=10)


def Skills_frame():
    global updater_frame
    
    if updater_frame is not None:
        Style(theme='superhero')
        updater_frame.config(bootstyle='info')#type:ignore
        # destroy any existing frames inside updater_frame
        for child in updater_frame.winfo_children():
            child.destroy()
            
        # create and add the Experience_frame
        Skills_frame = ttk.LabelFrame(updater_frame, text='Skills_frame', width=700, height=800,bootstyle='info') #type:ignore
        Skills_frame.pack(fill='both', padx=10, pady=10,ipady=100)
        
        skills_label = ttk.Label(Skills_frame,text='Skills',bootstyle='primary') #type:ignore
        skills_label.pack(side='top',anchor='w',padx=10,pady=10)
        skills_entry1 = ttk.Entry(Skills_frame,bootstyle='primary') #type:ignore
        skills_entry1.pack(side='top',anchor='w',padx=10,pady=10,fill=X)
        
        def insertData():
            global rrn_value
            
            Skills = skills_entry1.get()
            
            
            conn = sqlite3.connect('MainDB.db')

            # Create a cursor object to execute SQL commands
            c = conn.cursor()
            try:
                print(rrn_value,skills_entry1)
                c.execute('''UPDATE Resume
                            SET skills = ?
                            WHERE rrn_value = ?''',
                            (Skills, rrn_value))
                
                
                current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
                frame_name = "Skills"
                message = "Changes made in " + frame_name


                table.insert('', 'end', values=(current_time,message))


                conn.commit()
                messagebox.showinfo("Success", "Your Data Inserted Successfully!")
            except Exception as e:
                messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                conn.close()
        
        insert_data = tb.Button(Skills_frame,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
        insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

        save_button = tb.Button(Skills_frame,bootstyle='primary',text='Save') #type:ignore
        save_button.pack(side='left',anchor='se',pady=5,padx=10)

        Download_cv = tb.Button(Skills_frame,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
        Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
      

    else:
        # if updater_frame does not exist, create it and add the Skills_frame
        Updater_frame()
        Skills_frame = ttk.LabelFrame(updater_frame, text='Skills_frame', width=700, height=800)
        Skills_frame.pack(fill='both', padx=10, pady=10)


def project_frame():
    global updater_frame
    
    if updater_frame is not None:
        Style(theme='darkly')
        updater_frame.config(bootstyle='danger')#type:ignore
        # destroy any existing frames inside updater_frame
        for child in updater_frame.winfo_children():
            child.destroy()

        def projectframe2():
            global updater_frame
            
            if updater_frame is not None:
                Style(theme='darkly')
                for child in updater_frame.winfo_children():
                    child.destroy()
                project_frame = ttk.LabelFrame(updater_frame, text='project_frame2', width=700, height=800,bootstyle='danger') #type:ignore
                project_frame.pack(fill='both', padx=10, pady=10,ipady=300)
                
                project_name = ttk.Label(project_frame, text='Project Name',bootstyle='warning') #type:ignore
                project_name.pack(side=TOP, anchor='w' ,padx=10,pady=10)
                project_name_entry2 = ttk.Entry(project_frame)
                project_name_entry2.pack(side=TOP,anchor='w',padx=10,pady=10)
                
                explain_project = ttk.Label(project_frame,text='Explain Project',bootstyle='warning') #type:ignore
                explain_project.pack(side=TOP,anchor='w',padx=10,pady=10)
                explain_project_entry2 = ttk.Entry(project_frame) #type:ignore
                explain_project_entry2.pack(side=TOP,anchor='w',fill=X,padx=10,pady=10,ipady=20)
                
                def insertData():
                    global rrn_value
                    
                    project_name2 = project_name_entry2.get()
                    explain_project2 = explain_project_entry2.get()
                    
                    
                    conn = sqlite3.connect('MainDB.db')

                    # Create a cursor object to execute SQL commands
                    c = conn.cursor()
                    try:
                        print(rrn_value,project_name2,explain_project2)
                        c.execute('''UPDATE Resume
                                    SET project_name_two = ?, 
                                        explain_two_project = ?
                                    WHERE rrn_value = ?''',
                                    (project_name2, explain_project2, rrn_value))
                        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
                        frame_name = "Project Frame"
                        message = "Changes made in " + frame_name
                        table.insert('', 'end', values=(current_time,message))
                        conn.commit()
                        messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                    except Exception as e:
                        messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                        conn.close()
                
                insert_data = tb.Button(project_frame,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
                insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

                save_button = tb.Button(project_frame,bootstyle='primary',text='Save') #type:ignore
                save_button.pack(side='left',anchor='se',pady=5,padx=10)

                Download_cv = tb.Button(project_frame,bootstyle='primary',text='Download CV') #type:ignore
                Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
            
                Add_button = tb.Button(project_frame,bootstyle='primary',text='ADD',command=projectframe3) #type:ignore
                Add_button.pack(side='left',anchor='se',pady=5,padx=10)

                Next_button = tb.Button(project_frame,bootstyle='primary',text='Next',command=Experience) #type:ignore
                Next_button.pack(side='right',anchor='se',pady=5,padx=10)
        
        def projectframe3():
            global updater_frame
            
            if updater_frame is not None:
                Style(theme='darkly')
                for child in updater_frame.winfo_children():
                    child.destroy()
                project_frame = ttk.LabelFrame(updater_frame, text='project_frame3', width=700, height=800,bootstyle='danger') #type:ignore
                project_frame.pack(fill='both', padx=10, pady=10,ipady=300)
                
                project_name = ttk.Label(project_frame, text='Project Name',bootstyle='warning') #type:ignore
                project_name.pack(side=TOP, anchor='w' ,padx=10,pady=10)
                project_name_entry3 = ttk.Entry(project_frame)
                project_name_entry3.pack(side=TOP,anchor='w',padx=10,pady=10)
                
                explain_project = ttk.Label(project_frame,text='Explain Project',bootstyle='warning') #type:ignore
                explain_project.pack(side=TOP,anchor='w',padx=10,pady=10)
                explain_project_entry3 = ttk.Entry(project_frame) #type:ignore
                explain_project_entry3.pack(side=TOP,anchor='w',fill=X,padx=10,pady=10,ipady=20)
                
                def insertData():
                    global rrn_value
                    
                    projectname3 = project_name_entry3.get()
                    explainproject3 = explain_project_entry3.get()
                    
                    
                    conn = sqlite3.connect('MainDB.db')

                    # Create a cursor object to execute SQL commands
                    c = conn.cursor()
                    try:
                        print(rrn_value,projectname3,explainproject3)
                        c.execute('''UPDATE Resume
                                    SET project_name_three = ?, 
                                        explain_three_project = ?
                                    WHERE rrn_value = ?''',
                                    (projectname3,explainproject3 , rrn_value))
                        current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        
                        frame_name = "Project Frame"
                        message = "Changes made in " + frame_name

                        # Insert the records into the table widget
                        table.insert('', 'end', values=(current_time,message))
                        conn.commit()
                        messagebox.showinfo("Success", "Your Data Inserted Successfully!")
                    except Exception as e:
                        messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                        conn.close()
                
                
                insert_data = tb.Button(project_frame,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
                insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

                save_button = tb.Button(project_frame,bootstyle='primary',text='Save') #type:ignore
                save_button.pack(side='left',anchor='se',pady=5,padx=10)

                Download_cv = tb.Button(project_frame,bootstyle='primary',text='Download CV') #type:ignore
                Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      

                Next_button = tb.Button(project_frame,bootstyle='primary',text='Next',command=Experience) #type:ignore
                Next_button.pack(side='right',anchor='se',pady=5,padx=10)
        
        
        # create and add the Experience_frame
        project_frame = ttk.LabelFrame(updater_frame, text='project_frame', width=700, height=800,bootstyle='danger') #type:ignore
        project_frame.pack(fill='both', padx=10, pady=10,ipady=300)
        
        project_name = ttk.Label(project_frame, text='Project Name',bootstyle='warning') #type:ignore
        project_name.pack(side=TOP, anchor='w' ,padx=10,pady=10)
        project_name_entry = ttk.Entry(project_frame)
        project_name_entry.pack(side=TOP,anchor='w',padx=10,pady=10)
        
        explain_project = ttk.Label(project_frame,text='Explain Project',bootstyle='warning') #type:ignore
        explain_project.pack(side=TOP,anchor='w',padx=10,pady=10)
        explain_project_entry = ttk.Entry(project_frame) #type:ignore
        explain_project_entry.pack(side=TOP,anchor='w',fill=X,padx=10,pady=10,ipady=20)
        
        def insertData():
            global rrn_value
            
            projectname = project_name_entry.get()
            explainproject = explain_project_entry.get()
            
            
            conn = sqlite3.connect('MainDB.db')

            # Create a cursor object to execute SQL commands
            c = conn.cursor()
            try:
                print(rrn_value,projectname,explainproject)
                c.execute('''UPDATE Resume
                            SET project_name_one = ?, 
                                explain_one_project = ?
                            WHERE rrn_value = ?''',
                            (projectname,explainproject , rrn_value))
                current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
                frame_name = "Project Frame"
                message = "Changes made in " + frame_name

                table.insert('', 'end', values=(current_time,message))
                conn.commit()
                messagebox.showinfo("Success", "Your Data Inserted Successfully!")
            except Exception as e:
                messagebox.showerror("Error", "There was an error inserting your data: " + str(e))
                conn.close()        
        
        insert_data = tb.Button(project_frame,bootstyle='primary',text='Insert Data',command=insertData) #type:ignore
        insert_data.pack(side='left',anchor='sw',pady=5,padx=10)  

        save_button = tb.Button(project_frame,bootstyle='primary',text='Save') #type:ignore
        save_button.pack(side='left',anchor='se',pady=5,padx=10)

        Download_cv = tb.Button(project_frame,bootstyle='primary',text='Download CV',command=Downloadcv) #type:ignore
        Download_cv.pack(side='left',anchor='s',pady=5,padx=10)      
      
        Add_button = tb.Button(project_frame,bootstyle='primary',text='ADD',command=projectframe2) #type:ignore
        Add_button.pack(side='left',anchor='se',pady=5,padx=10)

        Next_button = tb.Button(project_frame,bootstyle='primary',text='Next',command=Experience) #type:ignore
        Next_button.pack(side='right',anchor='se',pady=5,padx=10)
    else:
        # if updater_frame does not exist, create it and add the project_frame
        Updater_frame()
        project_frame = ttk.LabelFrame(updater_frame, text='project_frame', width=700, height=800)
        project_frame.pack(fill='both', padx=10, pady=10)
    
    



def buttons():
         
    #common nav
    Home_button = tb.Button(Nav_bar,text='HOME',command=Home) #type:ignore
    Home_button.pack(side='top',pady=10,padx=10,ipadx=26)

    Dashboard_button = tb.Button(Nav_bar,text='DASHBOARD',command=DashboardFrame) #type:ignore
    Dashboard_button.pack(side='top',pady=10,padx=10,ipadx=10)

    #updater nav
    Updater_list_frame = tb.LabelFrame(Nav_bar,text="Updater List")
    Updater_list_frame.pack(side='top',pady=10,padx=10,ipadx=10)

    Personal_details = tb.Button(Updater_list_frame,text='Personal Details',command=Personal_frame)
    Personal_details.pack(side='top',pady=10,padx=10,ipadx=10)

    Experience_and_internships = tb.Button(Updater_list_frame,text='Experience',command=Experience)
    Experience_and_internships.pack(side='top',pady=10,padx=10,ipadx=10)

    Education = tb.Button(Updater_list_frame,text='Education',command=Education_frame)
    Education.pack(side='top',pady=10,padx=10,ipadx=10)

    project_button = tb.Button(Updater_list_frame,text='Project',command=project_frame)
    project_button.pack(side='top',pady=10,padx=10,ipadx=10)

    skills = tb.Button(Updater_list_frame,text='Skills',command=Skills_frame)
    skills.pack(side='top',pady=10,padx=10,ipadx=10)

    #common nav
    Updater = tb.Button(Nav_bar,text='UPDATER',command=Updater_frame) #type:ignore
    Updater.pack(side='top',pady=10,padx=10,ipadx=20)

        
#common nav
Home_button = tb.Button(Nav_bar,text='HOME',command=Home) #type:ignore
Home_button.pack(side='top',pady=10,padx=10,ipadx=26)

Dashboard_button = tb.Button(Nav_bar, text='Dashboard', command=lambda: DashboardFrame(pdf_location="MIAN CONTENT.pdf"))
Dashboard_button.pack(side='top', pady=10, padx=10, ipadx=10)

#updater nav
Updater_list_frame = tb.LabelFrame(Nav_bar,text="Updater List")
Updater_list_frame.pack(side='top',pady=10,padx=10,ipadx=10)

Personal_details = tb.Button(Updater_list_frame,text='Personal Details',command=Personal_frame)
Personal_details.pack(side='top',pady=10,padx=10,ipadx=10)

Experience_and_internships = tb.Button(Updater_list_frame,text='Experience',command=Experience)
Experience_and_internships.pack(side='top',pady=10,padx=10,ipadx=10)

Education = tb.Button(Updater_list_frame,text='Education',command=Education_frame)
Education.pack(side='top',pady=10,padx=10,ipadx=10)

project_button = tb.Button(Updater_list_frame,text='Project',command=project_frame)
project_button.pack(side='top',pady=10,padx=10,ipadx=10)

skills = tb.Button(Updater_list_frame,text='Skills',command=Skills_frame)
skills.pack(side='top',pady=10,padx=10,ipadx=10)



#common nav
Updater = tb.Button(Nav_bar,text='UPDATER',command=Updater_frame) #type:ignore
Updater.pack(side='top',pady=10,padx=10,ipadx=20)



root.mainloop()





